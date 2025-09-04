# src/core/universal_parser.py
import re
from pathlib import Path
from typing import Optional, List
import mimetypes
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import pytesseract
from PIL import Image
# from sentence_transformers import SentenceTransformer, util  # Sentence-BERT (暂时禁用，解决sqlite3 DLL问题)
from src.utils.logger import setup_logger

logger = setup_logger()

class Document:
    """文档对象：内容和元数据"""
    def __init__(self, content: str, metadata: dict = None):
        self.content = content
        self.metadata = metadata or {}

class UniversalParser:
    """多格式文档解析器（基于文档3.2.1，增强版）"""

    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'text/markdown': 'md',
        'image/jpeg': 'image',
        'image/png': 'image',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        # 添加更多如code（简单文本处理）
    }

    def __init__(self):
        # 暂时禁用SentenceTransformer，解决sqlite3 DLL问题
        self.sentence_model = None  # 原: SentenceTransformer('all-MiniLM-L6-v2')

    def detect_mime_type(self, file_path: str) -> str:
        """检测文件MIME类型"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or 'application/octet-stream'

    def parse(self, file_path: str) -> Optional[Document]:
        """解析文件内容"""
        path = Path(file_path)
        if not path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None
        
        mime_type = self.detect_mime_type(file_path)
        parser_method = self.SUPPORTED_MIME_TYPES.get(mime_type)
        
        if not parser_method:
            logger.warning(f"不支持的MIME类型: {mime_type} for {file_path}")
            return self._fallback_parse(file_path)
        
        try:
            content = getattr(self, f'_parse_{parser_method}')(file_path)
            processed_content = self.post_process(content, file_path)
            chunks = self.semantic_chunk(processed_content)  # 语义分块
            return Document('\n\n'.join(chunks), {'file_path': file_path, 'mime_type': mime_type})
        except Exception as e:
            logger.error(f"解析失败 {file_path}: {e}")
            return None

    def _parse_pdf(self, file_path: str) -> str:
        """PDF解析（使用PyPDF2）"""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return ' '.join(page.extract_text() for page in reader.pages if page.extract_text())

    def _parse_docx(self, file_path: str) -> str:
        """DOCX解析"""
        doc = DocxDocument(file_path)
        return '\n'.join(para.text for para in doc.paragraphs)

    def _parse_txt(self, file_path: str) -> str:
        """TXT/MD解析"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _parse_md(self, file_path: str) -> str:
        return self._parse_txt(file_path)  # MD作为TXT处理，保留格式

    def _parse_image(self, file_path: str) -> str:
        """图像OCR解析（集成pytesseract）"""
        img = Image.open(file_path)
        return pytesseract.image_to_string(img, lang='eng+chi_sim')  # 支持中英

    def _parse_xlsx(self, file_path: str) -> str:
        """XLSX解析（保留表格结构）"""
        df = pd.read_excel(file_path)
        return df.to_markdown(index=False)  # 转换为Markdown表格

    def _fallback_parse(self, file_path: str) -> Optional[Document]:
        """回退解析（简单文本）"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return Document(f.read())
        except UnicodeDecodeError:
            logger.warning(f"回退解析失败（非文本）: {file_path}")
            return None

    def post_process(self, content: str, file_path: str) -> str:
        """后处理：移除页眉/页脚，保留表格/代码（基于文档）"""
        # 移除常见页眉/页脚模式
        content = re.sub(r'Page \d+ of \d+', '', content)
        content = re.sub(r'^\s*Confidential\s*$', '', content, flags=re.MULTILINE)
        
        # 保留表格（假设已转换为MD）
        # 保留代码块（针对代码文件）
        if Path(file_path).suffix in ['.py', '.js', '.java']:
            content = self._preserve_code_blocks(content)
        
        return content.strip()

    def _preserve_code_blocks(self, content: str) -> str:
        """保留代码块格式"""
        lines = content.split('\n')
        in_code = False
        preserved = []
        for line in lines:
            if line.strip().startswith('def ') or line.strip().startswith('function '):
                in_code = True
            if in_code:
                preserved.append('    ' + line)  # 缩进保留
            else:
                preserved.append(line)
            if in_code and line.strip() == '':
                in_code = False
        return '\n'.join(preserved)

    def semantic_chunk(self, content: str, chunk_size: int = 512, overlap: int = 128) -> List[str]:
        """简化版分块引擎：在SentenceTransformer不可用时提供基础功能"""
        # 基础版本：简单基于字符数分块，不使用语义相似度
        chunks = []
        current_chunk = []
        current_length = 0
        
        # 简单地按句子分割
        sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', content)
        
        for sent in sentences:
            sent_len = len(sent)
            if current_length + sent_len > chunk_size:
                # 到达大小限制，创建新块
                chunks.append(' '.join(current_chunk))
                # 实现重叠（如果设置了）
                if overlap > 0 and current_chunk:
                    # 计算重叠句子数量
                    overlap_chars = 0
                    overlap_sents = []
                    for i in range(len(current_chunk)-1, -1, -1):
                        overlap_sents.insert(0, current_chunk[i])
                        overlap_chars += len(current_chunk[i]) + 1  # +1 for space
                        if overlap_chars >= overlap:
                            break
                    current_chunk = overlap_sents
                else:
                    current_chunk = []
                current_length = sum(len(s) for s in current_chunk) + 1  # +1 for space
            
            current_chunk.append(sent)
            current_length += sent_len + 1  # +1 for space
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        # 保留表格/代码完整性：如果chunk包含表格/代码标记，确保不拆分
        for i in range(len(chunks) - 1):
            if '|' in chunks[i] and '|' in chunks[i+1]:  # 简单表格检测
                chunks[i] += '\n' + chunks[i+1]
                chunks[i+1] = ''
        
        return [c for c in chunks if c]  # 移除空chunk