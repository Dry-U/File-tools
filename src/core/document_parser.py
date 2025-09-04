#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""文档解析器模块 - 处理多种格式文档的内容提取"""
import os
import time
import logging
import PyPDF2
import pdfminer.high_level
import pdfminer.layout
from docx import Document as DocxDocument
import pandas as pd
import markdown
import textract
from PIL import Image
import exifread
import datetime

class DocumentParser:
    """文档解析器类，用于提取各种格式文档的内容和元数据"""
    def __init__(self, config):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # 支持的文件类型映射到对应的解析函数
        self.parser_map = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'txt': self._parse_text,
            'md': self._parse_markdown,
            'csv': self._parse_csv,
            'xlsx': self._parse_excel,
            'xls': self._parse_excel,
            'jpg': self._parse_image,
            'jpeg': self._parse_image,
            'png': self._parse_image,
            'gif': self._parse_image,
            # 代码文件
            'py': self._parse_text,
            'java': self._parse_text,
            'cpp': self._parse_text,
            'js': self._parse_text,
            'html': self._parse_text,
            'css': self._parse_text,
        }
    
    def extract_text(self, file_path):
        """提取文件文本内容"""
        if not os.path.exists(file_path):
            self.logger.error(f"文件不存在: {file_path}")
            return f"错误: 文件不存在"
        
        file_ext = os.path.splitext(file_path)[1].lower()[1:]  # 获取文件扩展名，去除点号
        
        try:
            # 尝试使用特定的解析器
            if file_ext in self.parser_map:
                return self.parser_map[file_ext](file_path)
            else:
                # 尝试使用通用解析器作为后备
                return self._parse_generic(file_path)
        except Exception as e:
            self.logger.error(f"解析文件失败 {file_path}: {str(e)}")
            return f"错误: 无法解析文件内容\n{str(e)}"
    
    def extract_metadata(self, file_path):
        """提取文件元数据"""
        if not os.path.exists(file_path):
            self.logger.error(f"文件不存在: {file_path}")
            return {"错误": "文件不存在"}
        
        metadata = {}
        
        try:
            # 获取基本文件信息
            stat_info = os.stat(file_path)
            metadata["文件路径"] = file_path
            metadata["文件名"] = os.path.basename(file_path)
            metadata["文件大小"] = f"{stat_info.st_size / 1024:.2f} KB"
            metadata["创建时间"] = datetime.datetime.fromtimestamp(stat_info.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            metadata["修改时间"] = datetime.datetime.fromtimestamp(stat_info.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            metadata["访问时间"] = datetime.datetime.fromtimestamp(stat_info.st_atime).strftime('%Y-%m-%d %H:%M:%S')
            metadata["文件类型"] = os.path.splitext(file_path)[1].upper()
            
            # 获取特定格式的元数据
            file_ext = os.path.splitext(file_path)[1].lower()[1:]
            
            if file_ext == 'pdf':
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif file_ext == 'docx':
                docx_metadata = self._extract_docx_metadata(file_path)
                metadata.update(docx_metadata)
            elif file_ext in ['jpg', 'jpeg', 'png', 'gif']:
                image_metadata = self._extract_image_metadata(file_path)
                metadata.update(image_metadata)
            
        except Exception as e:
            self.logger.error(f"提取元数据失败 {file_path}: {str(e)}")
            metadata["元数据提取错误"] = str(e)
        
        return metadata
    
    def _parse_pdf(self, file_path):
        """解析PDF文件"""
        try:
            # 尝试使用PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() or ""
                
                # 如果PyPDF2提取的文本为空，尝试使用pdfminer
                if not text.strip():
                    text = pdfminer.high_level.extract_text(file_path)
                
                return text
        except Exception as e:
            self.logger.error(f"PDF解析失败 {file_path}: {str(e)}")
            # 尝试使用textract作为后备
            try:
                return textract.process(file_path).decode('utf-8', errors='ignore')
            except:
                return f"错误: 无法解析PDF内容\n{str(e)}"
    
    def _parse_docx(self, file_path):
        """解析Word文档"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            return text
        except Exception as e:
            self.logger.error(f"DOCX解析失败 {file_path}: {str(e)}")
            # 尝试使用textract作为后备
            try:
                return textract.process(file_path).decode('utf-8', errors='ignore')
            except:
                return f"错误: 无法解析Word内容\n{str(e)}"
    
    def _parse_text(self, file_path):
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk', errors='replace') as file:
                    return file.read()
            except:
                self.logger.error(f"文本解析失败 {file_path}: 编码错误")
                return "错误: 无法解析文本内容（编码问题）"
        except Exception as e:
            self.logger.error(f"文本解析失败 {file_path}: {str(e)}")
            return f"错误: 无法解析文本内容\n{str(e)}"
    
    def _parse_markdown(self, file_path):
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
                # 可以选择返回原始markdown或转换为HTML
                return content  # 返回原始markdown
        except Exception as e:
            self.logger.error(f"Markdown解析失败 {file_path}: {str(e)}")
            return f"错误: 无法解析Markdown内容\n{str(e)}"
    
    def _parse_csv(self, file_path):
        """解析CSV文件"""
        try:
            df = pd.read_csv(file_path, encoding_errors='replace')
            return df.to_string()
        except Exception as e:
            self.logger.error(f"CSV解析失败 {file_path}: {str(e)}")
            # 尝试作为普通文本文件解析
            try:
                return self._parse_text(file_path)
            except:
                return f"错误: 无法解析CSV内容\n{str(e)}"
    
    def _parse_excel(self, file_path):
        """解析Excel文件"""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            self.logger.error(f"Excel解析失败 {file_path}: {str(e)}")
            # 尝试使用textract作为后备
            try:
                return textract.process(file_path).decode('utf-8', errors='ignore')
            except:
                return f"错误: 无法解析Excel内容\n{str(e)}"
    
    def _parse_image(self, file_path):
        """解析图像文件"""
        try:
            # 对于图像，我们主要提取元数据而不是内容
            # 这里可以添加OCR功能，但需要额外的依赖
            with Image.open(file_path) as img:
                width, height = img.size
                mode = img.mode
                format_ = img.format
                
                content = f"图像文件信息:\n"
                content += f"- 尺寸: {width}x{height}\n"
                content += f"- 模式: {mode}\n"
                content += f"- 格式: {format_}\n"
                
                # 尝试提取EXIF信息
                try:
                    with open(file_path, 'rb') as f:
                        exif_data = exifread.process_file(f)
                        if exif_data:
                            content += "\nEXIF信息:\n"
                            # 只显示一些关键的EXIF信息
                            for tag in ['EXIF DateTimeOriginal', 'Image Make', 'Image Model', 'GPS GPSLatitude']:
                                if tag in exif_data:
                                    content += f"- {tag}: {exif_data[tag]}\n"
                except:
                    pass
                
                return content
        except Exception as e:
            self.logger.error(f"图像解析失败 {file_path}: {str(e)}")
            return f"错误: 无法解析图像内容\n{str(e)}"
    
    def _parse_generic(self, file_path):
        """通用解析器，用于处理不支持的文件格式"""
        try:
            # 使用textract尝试提取内容
            return textract.process(file_path).decode('utf-8', errors='ignore')
        except Exception as e:
            self.logger.error(f"通用解析失败 {file_path}: {str(e)}")
            return f"错误: 无法解析文件内容（不支持的格式）\n{str(e)}"
    
    def _extract_pdf_metadata(self, file_path):
        """提取PDF文件特定的元数据"""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                pdf_info = reader.metadata
                if pdf_info:
                    for key, value in pdf_info.items():
                        # 清理键名
                        clean_key = key.replace('/','')
                        metadata[f"PDF-{clean_key}"] = str(value)
                
                # 添加页数信息
                metadata["PDF-页数"] = len(reader.pages)
        except Exception as e:
            self.logger.error(f"提取PDF元数据失败 {file_path}: {str(e)}")
        
        return metadata
    
    def _extract_docx_metadata(self, file_path):
        """提取Word文档特定的元数据"""
        metadata = {}
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["Word-标题"] = core_props.title
            if core_props.author:
                metadata["Word-作者"] = core_props.author
            if core_props.subject:
                metadata["Word-主题"] = core_props.subject
            if core_props.keywords:
                metadata["Word-关键词"] = core_props.keywords
            if core_props.comments:
                metadata["Word-注释"] = core_props.comments
            if core_props.created:
                metadata["Word-创建时间"] = core_props.created.strftime('%Y-%m-%d %H:%M:%S')
            if core_props.modified:
                metadata["Word-修改时间"] = core_props.modified.strftime('%Y-%m-%d %H:%M:%S')
                
            # 计算段落数和页数
            metadata["Word-段落数"] = len(doc.paragraphs)
        except Exception as e:
            self.logger.error(f"提取Word元数据失败 {file_path}: {str(e)}")
        
        return metadata
    
    def _extract_image_metadata(self, file_path):
        """提取图像文件特定的元数据"""
        metadata = {}
        try:
            with Image.open(file_path) as img:
                metadata["图像-尺寸"] = f"{img.width}x{img.height}"
                metadata["图像-模式"] = img.mode
                metadata["图像-格式"] = img.format
                metadata["图像-颜色数"] = len(img.getcolors(maxcolors=2**24)) if img.mode != 'RGB' else ' millions'
            
            # 尝试提取EXIF信息
            try:
                with open(file_path, 'rb') as f:
                    exif_data = exifread.process_file(f)
                    if exif_data:
                        # 提取一些常见的EXIF标签
                        exif_tags = {
                            'EXIF DateTimeOriginal': '拍摄时间',
                            'Image Make': '相机制造商',
                            'Image Model': '相机型号',
                            'EXIF FNumber': '光圈',
                            'EXIF ExposureTime': '曝光时间',
                            'EXIF ISOSpeedRatings': 'ISO',
                            'EXIF FocalLength': '焦距'
                        }
                        
                        for tag, label in exif_tags.items():
                            if tag in exif_data:
                                metadata[f"图像-{label}"] = str(exif_data[tag])
            except Exception as e:
                self.logger.warning(f"提取EXIF信息失败 {file_path}: {str(e)}")
        except Exception as e:
            self.logger.error(f"提取图像元数据失败 {file_path}: {str(e)}")
        
        return metadata